from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import Workbook, load_workbook
except ImportError:
    Workbook = None
    load_workbook = None


SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".xlsm"}


@dataclass
class ExtractedDocument:
    source: Path
    text: str
    tables: list[list[list[str]]] = field(default_factory=list)


def normalize_text(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_docx(path: Path) -> ExtractedDocument:
    if Document is None:
        raise RuntimeError(
            "读取 DOCX 需要 python-docx。请运行：python -m pip install python-docx"
        )
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    table_lines: list[str] = []

    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
                table_lines.append(" | ".join(cells))
        if rows:
            tables.append(rows)
            table_lines.append("")

    joined = "\n".join(paragraphs)
    if table_lines:
        joined += "\n\n## 表格抽取\n" + "\n".join(table_lines)
    return ExtractedDocument(source=path, text=normalize_text(joined), tables=tables)


def read_xlsx(path: Path) -> ExtractedDocument:
    if load_workbook is None:
        raise RuntimeError(
            "读取 Excel 需要 openpyxl。请运行：python -m pip install openpyxl"
        )
    wb = load_workbook(path, data_only=True, read_only=True)
    lines: list[str] = []
    tables: list[list[list[str]]] = []

    for sheet in wb.worksheets:
        lines.append(f"## 工作表：{sheet.title}")
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if value is None else str(value).strip() for value in row]
            if any(cells):
                rows.append(cells)
                lines.append(" | ".join(cells))
        if rows:
            tables.append(rows)
        lines.append("")

    return ExtractedDocument(source=path, text=normalize_text("\n".join(lines)), tables=tables)


def read_pdf(path: Path) -> ExtractedDocument:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("缺少 pdfplumber，暂时无法读取 PDF 文本。") from exc

    lines: list[str] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                lines.append(f"## 第 {index} 页")
                lines.append(text.strip())
                lines.append("")
    text = normalize_text("\n".join(lines))
    if not text:
        raise RuntimeError(
            "PDF 中未提取到文本，文件可能是扫描件。请先执行 OCR，再审查 OCR 结果。"
        )
    return ExtractedDocument(source=path, text=text)


def read_document(path: Path) -> ExtractedDocument:
    if not path.exists():
        raise FileNotFoundError(f"输入文件不存在：{path}")
    if not path.is_file():
        raise ValueError(f"输入路径不是文件：{path}")

    suffix = path.suffix.lower()
    if suffix == ".docx":
        return read_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx(path)
    if suffix == ".pdf":
        return read_pdf(path)
    supported = "、".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(f"暂不支持的文件类型：{path.suffix}；当前支持 {supported}")


def first_match(text: str, pattern: str, default: str = "") -> str:
    match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
    if not match:
        return default
    value = match.group(1) if match.groups() else match.group(0)
    return clean_value(value)


def label_value(text: str, label_pattern: str) -> str:
    match = re.search(rf"^{label_pattern}[ \t]*[:：][ \t]*(.*)$", text, re.MULTILINE)
    if not match:
        return ""
    return clean_value(match.group(1))


def all_matches(text: str, pattern: str) -> list[str]:
    values = []
    for match in re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE):
        value = match.group(1) if match.groups() else match.group(0)
        value = clean_value(value)
        if value and value not in values:
            values.append(value)
    return values


def extract_contacts(text: str) -> list[dict[str, str]]:
    contacts: list[dict[str, str]] = []
    patterns = [
        re.compile(r"([\u4e00-\u9fff]{2,4})[ \t：:]*(1[3-9]\d{9})"),
        re.compile(r"(1[3-9]\d{9})[ \t：:]*(?:、|，|,)?[ \t]*([\u4e00-\u9fff]{2,4})"),
    ]
    seen: set[tuple[str, str]] = set()
    role_terms = ["项目经理", "安全员", "技术负责人", "施工负责人", "联系人"]

    for pattern_index, pattern in enumerate(patterns):
        for match in pattern.finditer(text):
            if pattern_index == 0:
                name, phone = match.group(1), match.group(2)
            else:
                phone, name = match.group(1), match.group(2)
            key = (clean_value(name), phone)
            if key in seen:
                continue
            seen.add(key)
            window = text[max(0, match.start() - 40) : match.end() + 40]
            nearby_roles = [role for role in role_terms if role in window]
            contacts.append(
                {
                    "name": key[0],
                    "phone": phone,
                    "role_note": "、".join(nearby_roles),
                }
            )
    return contacts


def clean_value(value: str) -> str:
    value = value.strip()
    value = re.sub(r"\s+", " ", value)
    value = value.strip(" ：:，,。；;")
    return value


def split_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[。；;])\s*|\n+", text)
    return [clean_value(part) for part in parts if clean_value(part)]


def detect_document_type(text: str) -> str:
    checks = [
        ("cryogenic_oxygen_installation", ["深冷", "制氧", "氧气", "禁油"]),
        ("steel_structure_installation", ["钢结构", "高强螺栓", "钢柱", "钢梁", "吊装"]),
        ("fire_control_merger", ["消防中控室", "火灾自动报警", "消防控制室", "CAN总线"]),
    ]
    scores = {
        doc_type: sum(text.count(keyword) for keyword in keywords)
        for doc_type, keywords in checks
    }
    doc_type, score = max(scores.items(), key=lambda item: item[1])
    return doc_type if score > 0 else "engineering_document"


def detect_document_title(lines: list[str]) -> str:
    for line in lines[:20]:
        if "方案" in line or "施工组织设计" in line:
            return line
    return lines[1] if len(lines) > 1 else ""


def extract_outline(lines: list[str], limit: int = 80) -> list[str]:
    outline: list[str] = []
    heading_pattern = re.compile(
        r"^("
        r"[一二三四五六七八九十]+[、.．]\s*.+|"
        r"（[一二三四五六七八九十]+）.+|"
        r"\d+(?:\.\d+){0,3}\s+.+|"
        r"附件\d+[:：].+"
        r")$"
    )
    for line in lines:
        if heading_pattern.match(line) and len(line) <= 80:
            if line not in outline:
                outline.append(line)
        if len(outline) >= limit:
            break
    return outline


def extract_key_quantities(text: str, limit: int = 40) -> list[str]:
    pattern = re.compile(
        r"[^。\n；;]{0,24}"
        r"\d[\d,]*(?:\.\d+)?\s*(?:日历天|天|小时|h|m²|㎡|m³|Nm³/h|MPa|℃|°C|米|m|台|套|个|只|块|路|点|根|榀|组|t|吨|KPa|kPa|MΩ|Ω)"
        r"[^。\n；;]{0,30}"
    )
    values: list[str] = []
    for match in pattern.finditer(text):
        value = clean_value(match.group(0))
        if value and value not in values:
            values.append(value)
        if len(values) >= limit:
            break
    return values


def table_header(row: list[str]) -> str:
    return "|".join(clean_value(cell) for cell in row if clean_value(cell))


def find_tables(doc: ExtractedDocument, required_keywords: list[str], limit: int = 5) -> list[dict[str, Any]]:
    matches: list[dict[str, Any]] = []
    for index, table in enumerate(doc.tables, start=1):
        if not table:
            continue
        header = table_header(table[0])
        table_text = "\n".join(" | ".join(row) for row in table[:10])
        haystack = header + "\n" + table_text
        if all(keyword in haystack for keyword in required_keywords):
            matches.append(
                {
                    "table_index": index,
                    "header": table[0],
                    "rows": table[1:],
                }
            )
        if len(matches) >= limit:
            break
    return matches


def snippets_for(text: str, keywords: list[str], limit: int = 8) -> list[str]:
    snippets: list[str] = []
    for sentence in split_sentences(text):
        if any(keyword in sentence for keyword in keywords):
            if sentence not in snippets:
                snippets.append(sentence)
        if len(snippets) >= limit:
            break
    return snippets


def extract_stages(text: str) -> list[dict[str, str]]:
    stages: list[dict[str, str]] = []
    pattern = re.compile(r"(第[一二三四五六七八九十]+阶段)：([^（\n]+)（(\d+)天）→?\s*([^\n]*)")
    for match in pattern.finditer(text):
        stages.append(
            {
                "stage": clean_value(match.group(1)),
                "name": clean_value(match.group(2)),
                "duration_days": match.group(3),
                "detail": clean_value(match.group(4)),
            }
        )
    return stages


def extract_civil_schedule(text: str) -> list[dict[str, str]]:
    paragraph = first_match(text, r"(土建工程15天施工节点计划：.+?)(?:\n|第三阶段|$)")
    if not paragraph:
        return []

    items: list[dict[str, str]] = []
    pattern = re.compile(r"第(\d+)(?:-(\d+))?天([^；。]+)")
    for match in pattern.finditer(paragraph):
        start = match.group(1)
        end = match.group(2) or start
        items.append(
            {
                "day": f"{start}-{end}" if start != end else start,
                "task": clean_value(match.group(3)),
            }
        )
    return items


def extract_schedule_from_tables(doc: ExtractedDocument) -> list[dict[str, Any]]:
    schedules: list[dict[str, Any]] = []
    for table in find_tables(doc, ["阶段", "工期"], limit=10):
        rows = []
        for row in table["rows"]:
            if any(clean_value(cell) for cell in row):
                rows.append(row)
        schedules.append(
            {
                "table_index": table["table_index"],
                "header": table["header"],
                "rows": rows,
            }
        )
    return schedules


def extract_type_specific(doc_type: str, doc: ExtractedDocument) -> dict[str, Any]:
    text = doc.text
    data: dict[str, Any] = {
        "key_quantities": extract_key_quantities(text),
        "equipment_or_material_tables": find_tables(doc, ["名称", "数量"], limit=8),
        "risk_tables": find_tables(doc, ["风险", "措施"], limit=5),
        "schedule_tables": extract_schedule_from_tables(doc),
    }

    if doc_type == "fire_control_merger":
        data["fire_control"] = {
            "total_alarm_points": first_match(text, r"三个品牌合计原有主机\d+台，监控报警点位([\d,]+个)"),
            "brands": {
                "lida_points": first_match(text, r"利达品牌[^。]*?合计([\d,]+个报警点位)"),
                "jbf_points": first_match(text, r"青鸟品牌[^。]*?合计([\d,]+个报警点位)"),
                "gt_points": first_match(text, r"国泰怡安品牌[^。]*?管理([\d,]+个报警点位)"),
            },
            "cable": first_match(text, r"(30×1\.5mm²电缆[\d,]+米)"),
            "fiber": first_match(text, r"(单模光纤[\d,]+米)"),
            "constraints": snippets_for(text, ["24小时连续运行", "不得整体停用", "分品牌独立", "消防通道"], limit=8),
        }
    elif doc_type == "steel_structure_installation":
        data["steel_structure"] = {
            "building_area": first_match(text, r"建筑面积\s*\|\s*([^\n|]+)") or first_match(text, r"建筑面积[:：]?\s*([\d,.]+㎡)"),
            "building_height": first_match(text, r"建筑高度\s*\|\s*([^\n|]+)") or first_match(text, r"建筑高度[:：]?\s*([^\n，。]+)"),
            "structure_type": first_match(text, r"结构形式\s*\|\s*([^\n|]+)") or first_match(text, r"结构形式[:：]?\s*([^\n，。]+)"),
            "max_span": first_match(text, r"钢框架梁最大跨度[:：]\s*([^\n]+)"),
            "max_single_weight": first_match(text, r"最大单件重量[:：]\s*([^\n]+)"),
            "max_install_height": first_match(text, r"最大安装高度[:：]\s*([^\n]+)"),
            "process_flow": snippets_for(text, ["总体施工流程", "先柱后梁", "先主后次", "逐层推进"], limit=10),
        }
    elif doc_type == "cryogenic_oxygen_installation":
        data["cryogenic_oxygen"] = {
            "oxygen_output": first_match(text, r"产品氧气产量\s*([0-9]+\s*Nm³/h)"),
            "nitrogen_output": first_match(text, r"产品氮气产量\s*([0-9]+\s*Nm³/h)"),
            "low_temperature": first_match(text, r"最低\s*([-\u2212]?\d+°C)"),
            "special_controls": snippets_for(text, ["禁油", "脱脂", "深冷", "氧气", "储罐"], limit=10),
        }
    return data


def schedule_claims(text: str, label: str) -> list[dict[str, str]]:
    claims: list[dict[str, str]] = []
    pattern = re.compile(
        rf"[^。\n；;]{{0,45}}{label}[^。\n；;]{{0,30}}?(\d+)\s*(日历天|天)[^。\n；;]{{0,30}}"
    )
    seen: set[tuple[str, str]] = set()
    for match in pattern.finditer(text):
        value = f"{match.group(1)}{match.group(2)}"
        context = clean_value(match.group(0))
        key = (value, context)
        if key not in seen:
            seen.add(key)
            claims.append({"value": value, "context": context})
    return claims


def build_review_checks(fields: dict[str, Any], text: str) -> dict[str, list[Any]]:
    basic = fields["basic"]
    schedule = fields["schedule"]
    safety = fields["safety"]
    missing: list[str] = []
    conflicts: list[dict[str, Any]] = []
    warnings: list[str] = []

    if not basic.get("project_name"):
        missing.append("未识别到项目名称。")
    if not basic.get("document_title"):
        missing.append("未识别到文档标题。")
    if (
        not schedule.get("total_duration")
        and not schedule.get("civil_schedule")
        and not schedule.get("schedule_tables")
    ):
        missing.append("未识别到施工进度计划。")
    if not safety.get("key_snippets"):
        missing.append("未识别到安全风险或应急措施摘录。")

    for label, display_name in [("总工期", "总工期"), ("土建(?:工程)?工期", "土建工期")]:
        claims = schedule_claims(text, label)
        unique_values = sorted({claim["value"] for claim in claims})
        if len(unique_values) > 1:
            conflicts.append(
                {
                    "field": display_name,
                    "values": unique_values,
                    "evidence": claims,
                    "note": "同一类工期出现多个值，需要人工确认是否为修订前后版本或不同范围。",
                }
            )

    if fields["contacts"]["phone_numbers"] and not fields["contacts"]["people"]:
        warnings.append("识别到电话号码，但未能可靠关联姓名或岗位。")
    if fields["document_stats"]["characters"] < 200:
        warnings.append("可提取文本较少，结果可能不完整；扫描件应先做 OCR。")
    warnings.append("本工具只做本地预审，不判断法规符合性，也不替代注册工程师或项目负责人审核。")

    return {
        "missing_items": missing,
        "conflicts": conflicts,
        "warnings": warnings,
    }


def extract_fields(doc: ExtractedDocument) -> dict[str, Any]:
    text = doc.text
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    doc_type = detect_document_type(text)
    project_name = (
        label_value(text, r"项目名称")
        or first_match(text, r"^(KDON[-/0-9]+[^\n]*项目)")
        or (lines[0] if lines else "")
    )
    document_title = detect_document_title(lines)

    fields: dict[str, Any] = {
        "source_file": doc.source.name,
        "extracted_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "document_type": doc_type,
        "basic": {
            "project_name": project_name,
            "document_title": document_title,
            "project_code": label_value(text, r"项目代号") or label_value(text, r"工程编号") or label_value(text, r"申报编号"),
            "construction_unit": label_value(text, r"建设单位"),
            "design_unit": label_value(text, r"设计单位"),
            "document_date": label_value(text, r"日\s*期"),
            "prepared_by": label_value(text, r"编\s*制\s*人"),
            "reviewed_by": label_value(text, r"审\s*核\s*人"),
            "approved_by": label_value(text, r"批\s*准\s*人"),
            "compiler_unit": label_value(text, r"编制单位"),
            "location": first_match(text, r"位于([^。；\n]+)"),
        },
        "contacts": {
            "phone_numbers": all_matches(text, r"(?<!\d)(1[3-9]\d{9})(?!\d)"),
            "people": extract_contacts(text),
        },
        "scope": {
            "equipment_count": first_match(text, r"共\s*([0-9]+类[0-9]+台/套设备)"),
            "equipment_total_weight": first_match(text, r"设备已知总重约\s*([0-9.]+吨)"),
            "pipeline_total": first_match(text, r"合计\s*([0-9]+米)"),
            "carbon_steel_pipe": first_match(text, r"碳钢管（[^）]+）\s*([0-9]+米)"),
            "stainless_pipe": first_match(text, r"不锈钢管（[^）]+）\s*([0-9]+米)"),
            "main_work_items": snippets_for(text, ["土建工程", "设备安装", "工艺管道", "电气仪表", "防腐绝热"], limit=8),
        },
        "process_parameters": {
            "inlet_air": first_match(text, r"入塔空气量\s*([0-9]+\s*Nm³/h)"),
            "oxygen_output": first_match(text, r"产品氧气产量\s*([0-9]+\s*Nm³/h)"),
            "nitrogen_output": first_match(text, r"产品氮气产量\s*([0-9]+\s*Nm³/h)"),
            "oxygen_pressure": first_match(text, r"氮气/氧气系统\s*≤\s*([0-9.]+\s*MPa)"),
            "low_temperature": first_match(text, r"最低\s*([-\u2212]?\d+°C)"),
        },
        "schedule": {
            "total_duration": (
                first_match(text, r"施工总工期相应调整为约\s*([0-9]+天)")
                or first_match(text, r"总工期计划\s*([0-9]+日历天)")
                or first_match(text, r"钢结构安装总工期计划\s*([0-9]+日历天)")
                or first_match(text, r"项目总工期[^0-9]{0,20}([0-9]+日历天|[0-9]+天)")
            ),
            "civil_duration": first_match(text, r"土建工程工期调整为\s*([0-9]+天)"),
            "stages": extract_stages(text),
            "civil_schedule": extract_civil_schedule(text),
            "schedule_tables": extract_schedule_from_tables(doc),
        },
        "safety": {
            "risk_keywords": {
                "储罐": text.count("储罐"),
                "氧气": text.count("氧气"),
                "禁油": text.count("禁油"),
                "火源": text.count("火源"),
                "应急": text.count("应急"),
                "120": text.count("120"),
                "吊装": text.count("吊装"),
                "深冷": text.count("深冷"),
                "特种设备": text.count("特种设备"),
            },
            "key_snippets": snippets_for(
                text,
                ["储罐", "氧气", "禁油", "火源", "应急", "120", "吊装", "深冷", "特种设备"],
                limit=12,
            ),
        },
        "document_stats": {
            "characters": len(text),
            "paragraph_like_lines": len(lines),
            "tables": len(doc.tables),
        },
        "outline": extract_outline(lines),
        "type_specific": extract_type_specific(doc_type, doc),
        "pending_questions": [],
    }

    if not fields["contacts"]["phone_numbers"]:
        fields["pending_questions"].append("未识别到联系人电话。")
    if not fields["schedule"]["total_duration"] and not fields["schedule"]["civil_schedule"] and not fields["schedule"]["schedule_tables"]:
        fields["pending_questions"].append("未识别到施工进度计划。")
    if fields["safety"]["risk_keywords"].get("120", 0) == 0:
        fields["pending_questions"].append("未发现明确的120急救通知要求。")
    if doc_type == "cryogenic_oxygen_installation" and fields["safety"]["risk_keywords"].get("火源", 0) == 0:
        fields["pending_questions"].append("未发现明显的禁止火源说明。")

    fields["review_checks"] = build_review_checks(fields, text)
    for missing_item in fields["review_checks"]["missing_items"]:
        if missing_item not in fields["pending_questions"]:
            fields["pending_questions"].append(missing_item)

    return fields


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def write_excel(path: Path, fields: dict[str, Any]) -> None:
    if Workbook is None:
        raise RuntimeError(
            "生成 Excel 需要 openpyxl。请运行：python -m pip install openpyxl"
        )
    wb = Workbook()
    ws = wb.active
    ws.title = "字段"
    ws.append(["分组", "字段", "值"])

    def append_group(prefix: str, value: Any) -> None:
        if isinstance(value, dict):
            for key, child in value.items():
                append_group(f"{prefix}.{key}" if prefix else key, child)
        elif isinstance(value, list):
            ws.append([prefix, "", json.dumps(value, ensure_ascii=False)])
        else:
            group, _, key = prefix.rpartition(".")
            ws.append([group, key or prefix, value])

    append_group("", fields)
    wb.save(path)


def render_summary(fields: dict[str, Any]) -> str:
    basic = fields["basic"]
    contacts = fields["contacts"]
    scope = fields["scope"]
    schedule = fields["schedule"]
    safety = fields["safety"]

    lines: list[str] = []
    lines.append(f"# {basic.get('project_name') or '工程资料摘要'}")
    lines.append("")
    lines.append(f"> 文档类型：{fields.get('document_type', 'engineering_document')}")
    lines.append("")
    lines.append("## 基本信息")
    basic_rows = [
        ("文件", fields["source_file"]),
        ("文档", basic.get("document_title")),
        ("项目代号", basic.get("project_code")),
        ("建设单位", basic.get("construction_unit")),
        ("设计单位", basic.get("design_unit")),
        ("编制单位", basic.get("compiler_unit")),
        ("日期", basic.get("document_date")),
        ("位置", basic.get("location")),
    ]
    for label, value in basic_rows:
        if value:
            lines.append(f"- {label}：{value}")
    if contacts.get("people") or contacts.get("phone_numbers"):
        lines.append("")
        lines.append("## 联系人")
        linked_phones: set[str] = set()
        for person in contacts.get("people", []):
            linked_phones.add(person.get("phone", ""))
            role_note = (
                f"；附近岗位词：{person.get('role_note')}"
                if person.get("role_note")
                else ""
            )
            lines.append(
                f"- {person.get('name') or '姓名待确认'}：{person.get('phone')}{role_note}"
            )
        for phone in contacts.get("phone_numbers", []):
            if phone not in linked_phones:
                lines.append(f"- {phone}：姓名及岗位待确认")
    scope_summary = [
        ("设备数量", scope.get("equipment_count")),
        ("设备总重", scope.get("equipment_total_weight")),
        ("管道合计", scope.get("pipeline_total")),
        ("碳钢管", scope.get("carbon_steel_pipe")),
        ("不锈钢管", scope.get("stainless_pipe")),
    ]
    if any(value for _, value in scope_summary) or scope.get("main_work_items"):
        lines.append("")
        lines.append("## 工程范围与相关工作摘录")
        for label, value in scope_summary:
            if value:
                lines.append(f"- {label}：{value}")
    for item in scope.get("main_work_items", []):
        lines.append(f"- {item}")
    type_specific = fields.get("type_specific", {})
    key_quantities = type_specific.get("key_quantities", [])
    if key_quantities:
        lines.append("")
        lines.append("### 关键数量/参数")
        for item in key_quantities[:16]:
            lines.append(f"- {item}")
    lines.append("")
    lines.append("## 施工计划")
    if schedule.get("total_duration"):
        lines.append(f"- 总工期：{schedule.get('total_duration')}")
    if schedule.get("civil_duration"):
        lines.append(f"- 土建工期：{schedule.get('civil_duration')}")
    if not schedule.get("total_duration") and not schedule.get("civil_duration"):
        lines.append("- 未从正文中可靠识别总工期或土建工期。")
    if schedule.get("stages"):
        lines.append("")
        lines.append("| 阶段 | 名称 | 天数 | 摘要 |")
        lines.append("|---|---|---:|---|")
        for stage in schedule["stages"]:
            lines.append(f"| {stage['stage']} | {stage['name']} | {stage['duration_days']} | {stage['detail']} |")
    if schedule.get("civil_schedule"):
        lines.append("")
        lines.append("### 土建15天节点")
        for item in schedule["civil_schedule"]:
            lines.append(f"- 第{item['day']}天：{item['task']}")
    if schedule.get("schedule_tables"):
        lines.append("")
        lines.append("### 表格进度计划")
        for table in schedule["schedule_tables"][:2]:
            header = table.get("header", [])
            rows = table.get("rows", [])
            if not header or not rows:
                continue
            lines.append("")
            lines.append("| " + " | ".join(header) + " |")
            lines.append("|" + "|".join("---" for _ in header) + "|")
            for row in rows[:12]:
                padded = row[: len(header)] + [""] * max(0, len(header) - len(row))
                lines.append("| " + " | ".join(padded[: len(header)]) + " |")
    lines.append("")
    lines.append("## 安全与风险摘录")
    for key, count in safety.get("risk_keywords", {}).items():
        if count:
            lines.append(f"- {key}：出现 {count} 次")
    lines.append("")
    for item in safety.get("key_snippets", []):
        lines.append(f"> {item}")
        lines.append("")
    risk_tables = type_specific.get("risk_tables", [])
    if risk_tables:
        lines.append("### 风险表摘录")
        for table in risk_tables[:2]:
            header = table.get("header", [])
            rows = table.get("rows", [])
            if not header or not rows:
                continue
            lines.append("")
            lines.append("| " + " | ".join(header) + " |")
            lines.append("|" + "|".join("---" for _ in header) + "|")
            for row in rows[:10]:
                padded = row[: len(header)] + [""] * max(0, len(header) - len(row))
                lines.append("| " + " | ".join(padded[: len(header)]) + " |")
    if fields.get("outline"):
        lines.append("")
        lines.append("## 章节提纲")
        for item in fields["outline"][:40]:
            lines.append(f"- {item}")
    review_checks = fields.get("review_checks", {})
    lines.append("")
    lines.append("## 预审结果")
    missing_items = review_checks.get("missing_items") or ["未发现自动规则可判定的关键缺项。"]
    lines.append("### 缺项")
    for item in missing_items:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("### 冲突")
    conflicts = review_checks.get("conflicts") or []
    if not conflicts:
        lines.append("- 未发现自动规则可判定的字段冲突。")
    for item in conflicts:
        lines.append(
            f"- {item.get('field')}：{'、'.join(item.get('values', []))}。{item.get('note', '')}"
        )
        for evidence in item.get("evidence", []):
            lines.append(f"  - 证据：{evidence.get('context')}")
    lines.append("")
    lines.append("### 提醒")
    for item in review_checks.get("warnings", []):
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## 待确认问题")
    pending = fields.get("pending_questions") or ["暂无明显待确认项。"]
    for item in pending:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## 备注")
    lines.append("本报告由确定性规则在本地生成，正式交付或施工使用前必须人工复核。")
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="本地工程资料预审：提取字段、工期、安全风险、缺项与冲突。"
    )
    parser.add_argument("input", help="输入文件路径（DOCX、PDF、XLSX 或 XLSM）。")
    parser.add_argument(
        "--output-dir",
        help="输出目录。省略时输出到当前目录的 <文件名>_review_<时间戳>/。",
    )
    parser.add_argument(
        "--no-excel",
        action="store_true",
        help="不生成 fields.xlsx；缺少 openpyxl 时可用于处理 DOCX/PDF。",
    )
    parser.add_argument(
        "--skip-extracted-text",
        action="store_true",
        help="不保存 extracted_text.md。",
    )
    args = parser.parse_args()

    try:
        input_path = Path(args.input).expanduser().resolve()
        document = read_document(input_path)
        fields = extract_fields(document)

        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = (
            Path(args.output_dir).expanduser().resolve()
            if args.output_dir
            else Path.cwd() / f"{input_path.stem}_review_{stamp}"
        )
        output_dir.mkdir(parents=True, exist_ok=True)

        write_json(output_dir / "fields.json", fields)
        if not args.no_excel:
            write_excel(output_dir / "fields.xlsx", fields)
        (output_dir / "review_report.md").write_text(
            render_summary(fields), encoding="utf-8"
        )
        if not args.skip_extracted_text:
            (output_dir / "extracted_text.md").write_text(
                document.text, encoding="utf-8"
            )

        print(f"输入：{input_path}")
        print(f"输出：{output_dir}")
        print(f"- {output_dir / 'fields.json'}")
        if not args.no_excel:
            print(f"- {output_dir / 'fields.xlsx'}")
        print(f"- {output_dir / 'review_report.md'}")
        if not args.skip_extracted_text:
            print(f"- {output_dir / 'extracted_text.md'}")
    except (FileNotFoundError, RuntimeError, ValueError) as exc:
        parser.exit(2, f"错误：{exc}\n")


if __name__ == "__main__":
    main()
